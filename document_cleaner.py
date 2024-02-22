import re
from docx import Document

def remove_definitions_and_acronyms(acronym_file, target_file):
    # Load the documents
    acronym_doc = Document(acronym_file)
    target_doc = Document(target_file)

    # Create a dictionary of acronyms and definitions
    acronyms = {}
    for table in acronym_doc.tables:
        for row in table.rows:
            acronyms[row.cells[0].text] = row.cells[1].text

    # Iterate over the paragraphs in the target document
    for paragraph in target_doc.paragraphs:
        # Replace each instance of the definition with the acronym
        for acronym, definition in acronyms.items():
            paragraph.text = re.sub(r'\b' + re.escape(definition) + r'\b', '', paragraph.text, flags=re.IGNORECASE)
            # Remove parentheses surrounding the acronym
            paragraph.text = re.sub(r'\(' + re.escape(acronym) + r'\)', acronym, paragraph.text)
            # Remove extra spaces before the acronym
            paragraph.text = re.sub(r' +' + re.escape(acronym), ' ' + acronym, paragraph.text)

    # Prompt the user for the number of pages in the front matter
    front_matter_pages = int(input("Enter the number of pages in the front matter: "))

    # Skip the front matter pages
    paragraphs_to_process = target_doc.paragraphs[front_matter_pages:]

    # Iterate over the paragraphs in the target document
    for acronym, definition in acronyms.items():
        for paragraph in paragraphs_to_process:
            # Replace only the first instance of the acronym with the definition
            if acronym in paragraph.text:
                paragraph.text = re.sub(r'\b' + re.escape(acronym) + r'\b', definition + ' (' + acronym + ')', paragraph.text, count=1, flags=re.IGNORECASE)
                break

    # Save the modified document
    target_doc.save('modified_' + target_file)

if __name__ == "__main__":
    acronym_file = input("Enter the name of the docx file containing the table of acronyms: ")
    target_file = input("Enter the name of the target docx file to be manipulated: ")
    remove_definitions_and_acronyms(acronym_file, target_file)

