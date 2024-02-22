import re
import sys
from docx import Document

def find_acronyms(text):
    # Regular expression to find acronyms and their definitions
    acronyms = re.findall(r'([A-Za-z0-9\s]+)\s\(([A-Z]{2,})\)', text)
    return acronyms

def main(docx_file):
    # Open the docx file
    doc = Document(docx_file)

    # Initialize an empty dictionary to store acronyms and their definitions
    acronym_dict = {}

    # Iterate over each paragraph in the document
    for para in doc.paragraphs:
        # Find acronyms and their definitions in the paragraph
        acronyms = find_acronyms(para.text)
        for defn, acronym in acronyms:
            # If the acronym is not already in the dictionary, add it with its definition
            if acronym not in acronym_dict:
                # Split the definition by spaces and take the last number of words equal to the length of the acronym
                defn_words = defn.split()
                short_defn = ' '.join(defn_words[-len(acronym):])
                acronym_dict[acronym] = short_defn.strip()
                
    # Review the definitions with the user
    print("Review the definitions:")
    for acronym, definition in acronym_dict.items():
        user_input = input(f"Is this the correct definition for '{acronym}'? '{definition}' (Y/N): ")
        if user_input.lower() == 'n':
            corrected_defn = input(f"Please enter the correct definition for '{acronym}': ")
            acronym_dict[acronym] = corrected_defn.strip()

    # Create a new document for the acronyms table
    table_doc = Document()
    table = table_doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Acronym'
    hdr_cells[1].text = 'Definition'

    # Add acronyms and definitions to the table
    for acronym, defn in acronym_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = acronym
        row_cells[1].text = defn

    # Save the modified DOCX file
    modified_file_path = docx_file.replace('.docx', '_acronyms_table.docx')
    table_doc.save(modified_file_path)
    print(f"Modified DOCX file with acronyms table saved as: {modified_file_path}")

if __name__ == "__main__":
    # The first command-line argument is the docx file path
    docx_file = sys.argv[1]
    main(docx_file)

